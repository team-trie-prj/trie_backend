"""① 지능형 문서 텍스트 파싱 및 노이즈 필터링  (FNC-DAT-01)

PDF / DOCX (+ TXT) 를 순수 텍스트로 파싱하고, 검색 품질을 떨어뜨리는 노이즈를 정제한다.

노이즈 정제 전략:
  - 페이지 반복 머리말/꼬리말 제거 (다수 페이지에 공통 등장하는 라인)
  - 페이지 번호 단독 라인 제거 ("3", "- 3 -", "3 / 10" 등)
  - 하이픈 줄바꿈 결합 (word-\nword → wordword)
  - 제어문자 / 비가시문자 제거, 공백·개행 정규화
  - 의미 없는 기호 나열 라인 제거
"""

from __future__ import annotations

import os
import re
import unicodedata
from collections import Counter
from dataclasses import dataclass, field

SUPPORTED_EXTS = {".pdf", ".docx", ".txt", ".md"}

# 페이지 번호 형태의 단독 라인
_PAGE_NUM_RE = re.compile(r"^\s*[-–—]?\s*\d{1,4}\s*(?:/\s*\d{1,4})?\s*[-–—]?\s*$")
# 하이픈으로 끊긴 줄바꿈 (영문 위주)
_HYPHEN_BREAK_RE = re.compile(r"(\w)-\n(\w)")
# 3회 이상 연속 개행
_MULTI_NL_RE = re.compile(r"\n{3,}")
# 라인 내 2칸 이상 공백
_MULTI_SPACE_RE = re.compile(r"[ \t ]{2,}")
# 제어문자 (개행/탭 제외)
_CTRL_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")


@dataclass
class ParsedDocument:
    text: str
    doc_type: str
    title: str
    page_count: int = 0
    char_count: int = 0
    meta: dict = field(default_factory=dict)


# ---------------------------------------------------------------------------
# 포맷별 원문 추출 (페이지/문단 리스트)
# ---------------------------------------------------------------------------
def _extract_pdf(path: str) -> list[str]:
    import pdfplumber

    pages: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            pages.append(page.extract_text() or "")
    return pages


def _extract_docx(path: str) -> list[str]:
    import docx  # python-docx

    document = docx.Document(path)
    blocks: list[str] = [p.text for p in document.paragraphs]

    # 표 셀 텍스트도 수집 (지침서/기준서에 표가 많음)
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    return ["\n".join(blocks)]


def _extract_txt(path: str) -> list[str]:
    with open(path, encoding="utf-8", errors="ignore") as f:
        return [f.read()]


# ---------------------------------------------------------------------------
# 노이즈 정제
# ---------------------------------------------------------------------------
def _footer_key(line: str) -> str:
    """머리말/꼬리말 비교용 정규화 키 — 앞뒤 페이지 번호를 제거.

    '요구사항 명세서 1', '요구사항 명세서 2' → 모두 '요구사항 명세서' 로 수렴시켜
    페이지 번호가 달라도 동일 꼬리말로 인식한다.
    """
    s = line.strip()
    s = re.sub(r"[\s\-–—/]*\d+(?:\s*/\s*\d+)?[\s\-–—]*$", "", s)  # 뒤쪽 번호
    s = re.sub(r"^[\s\-–—/]*\d+[\s\-–—/]*", "", s)  # 앞쪽 번호
    return re.sub(r"\s+", " ", s).strip()


def _detect_repeated_lines(pages: list[str], min_ratio: float = 0.5) -> set[str]:
    """페이지 가장자리(첫/마지막 줄)에서 반복되는 머리말/꼬리말 키 탐지.

    가장자리로 한정해, 숫자로 끝나는 본문 라인이 오탐 제거되는 것을 방지한다.
    """
    if len(pages) < 3:
        return set()
    counter: Counter[str] = Counter()
    for page in pages:
        lines = [ln.strip() for ln in page.splitlines() if ln.strip()]
        if not lines:
            continue
        for edge in {lines[0], lines[-1]}:  # 첫/마지막 줄만 후보
            key = _footer_key(edge)
            if 2 <= len(key) <= 60:
                counter[key] += 1
    threshold = max(3, int(len(pages) * min_ratio))
    return {key for key, cnt in counter.items() if cnt >= threshold}


def _is_noise_line(line: str, repeated: set[str]) -> bool:
    stripped = line.strip()
    if not stripped:
        return False  # 빈 줄은 개행 정규화에서 처리
    if _PAGE_NUM_RE.match(stripped):
        return True
    key = _footer_key(stripped)
    if key and key in repeated:
        return True
    # 알파벳/숫자/한글이 하나도 없는 기호 나열 라인
    if not re.search(r"[0-9A-Za-z가-힣]", stripped):
        return True
    return False


def clean_text(pages: list[str]) -> str:
    """페이지 리스트 → 정제된 단일 텍스트."""
    repeated = _detect_repeated_lines(pages)

    cleaned_pages: list[str] = []
    for page in pages:
        page = _HYPHEN_BREAK_RE.sub(r"\1\2", page)
        page = _CTRL_RE.sub("", page)
        page = unicodedata.normalize("NFC", page)

        kept: list[str] = []
        for line in page.splitlines():
            if _is_noise_line(line, repeated):
                continue
            line = _MULTI_SPACE_RE.sub(" ", line).rstrip()
            kept.append(line)
        cleaned_pages.append("\n".join(kept))

    text = "\n\n".join(cp for cp in cleaned_pages if cp.strip())
    text = _MULTI_NL_RE.sub("\n\n", text)
    return text.strip()


# ---------------------------------------------------------------------------
# 진입점
# ---------------------------------------------------------------------------
def parse_document(path: str) -> ParsedDocument:
    ext = os.path.splitext(path)[1].lower()
    if ext not in SUPPORTED_EXTS:
        raise ValueError(f"지원하지 않는 문서 형식: {ext} (지원: {sorted(SUPPORTED_EXTS)})")

    if ext == ".pdf":
        pages = _extract_pdf(path)
        doc_type = "pdf"
    elif ext == ".docx":
        pages = _extract_docx(path)
        doc_type = "docx"
    else:
        pages = _extract_txt(path)
        doc_type = "txt" if ext == ".txt" else "md"

    text = clean_text(pages)
    title = os.path.splitext(os.path.basename(path))[0]

    return ParsedDocument(
        text=text,
        doc_type=doc_type,
        title=title,
        page_count=len(pages),
        char_count=len(text),
        meta={"source_path": path, "ext": ext},
    )
